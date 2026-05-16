"""
PT6A bom.docx → pt6a_bom.csv 结构化转换（不走 LLM）

理由：
  1. stage1 的 docx 分支会走 extract_tables + llm_to_csv（581 张表 × LLM 调用，慢且静默）
  2. PT6A BOM 是标准 IPC 6 列格式，结构规整，无需 LLM
  3. 论文 §3.2.2 BOM Agent 本来就定义为"读取结构化 Excel/ERP 导出文件"

输出列（与 stage1 _bom_df_to_entities_and_triples 兼容）：
  fig_item / part_id / part_name / nomenclature / category / qty / material / effect_code
"""
from pathlib import Path
import re
import sys

import pandas as pd
from docx import Document

SRC = Path("data/KG/PT6A/bom.docx")
DST = Path("data/KG/PT6A/bom.csv")


def classify_category(fig_item: str, nomenclature: str) -> str:
    """
    IPC 分类严规则（默认 Part，除非明确标识为装配体）：
      触发 Assembly 条件（任一满足）：
        1. nomenclature 含 ASSEMBLY / ASY / INSTALLATION / GROUP
        2. 以 'SET OF' / 'KIT' 起头（整组套件）
    其余一律为 Part。
    """
    nom_upper = nomenclature.upper()
    if not nom_upper:
        return "Part"
    if re.search(r"\bASSEMBLY\b|\bASY\b|\bINSTALLATION\b|\bGROUP\b", nom_upper):
        return "Assembly"
    if nom_upper.startswith(("SET OF ", "KIT ", "KIT,", "KIT OF ")):
        return "Assembly"
    return "Part"


def extract_rows_from_doc(doc: Document):
    """遍历所有表格，提取 BOM 行"""
    rows = []
    skipped_tables = 0
    bom_table_cnt = 0

    for ti, table in enumerate(doc.tables):
        if not table.rows:
            continue
        header_cells = [c.text.strip().upper() for c in table.rows[0].cells]
        header_joined = " | ".join(header_cells)

        # 识别 BOM 主表（含 PART NUMBER + NOMENCLATURE 列）
        if not ("PART NUMBER" in header_joined and "NOMENCLATURE" in header_joined):
            skipped_tables += 1
            continue
        bom_table_cnt += 1

        # 表格可能有 6 列标准或 7-8 列（含 REMARKS 等）
        # 列索引：假定前 6 列为标准 IPC
        for ri, row in enumerate(table.rows[1:], start=1):
            cells = [c.text.strip() for c in row.cells]
            if len(cells) < 5:
                continue
            fig_item = cells[0]
            part_num = cells[1] if len(cells) > 1 else ""
            # 部分表有 AIRLINE PART NUMBER 列，可跳过（index 2）
            nomenclature = cells[3] if len(cells) > 3 else ""
            effect_code = cells[4] if len(cells) > 4 else ""
            units = cells[5] if len(cells) > 5 else ""

            # 跳过整行空行
            if not (fig_item or part_num or nomenclature):
                continue
            # 跳过 nomenclature 为空的占位行（如 fig_item='1' 作为块开始标记）
            if not nomenclature.strip():
                continue
            # 跳过"EFF MODEL"行、"SUPPLIER CODE"等元信息
            if nomenclature.upper().startswith(("EFF MODEL", "SUPPLIER CODE")):
                continue
            # 跳过纯 dash / 分隔符行
            if re.match(r"^[-\s\.]+$", nomenclature):
                continue
            # nomenclature 清洗：去除尾部的 SUPPLIER CODE
            nom_clean = re.sub(r"\s*SUPPLIER CODE[:\s].*$", "", nomenclature, flags=re.IGNORECASE).strip()
            # 生成 part_name：取前 60 字符
            part_name = nom_clean.split("1234567")[0].strip()[:80] if "1234567" in nom_clean else nom_clean[:80]
            if not part_name:
                part_name = nom_clean[:80] or fig_item
            # part_id：优先 part_num，否则用 fig_item
            pid = part_num.strip() if part_num.strip() else f"FIG-{fig_item}-{ti}-{ri}"

            rows.append({
                "fig_item": fig_item,
                "part_id": pid,
                "part_name": part_name,
                "nomenclature": nom_clean,
                "category": classify_category(fig_item, nom_clean),
                "qty": units or "1",
                "material": "",
                "effect_code": effect_code,
            })

    return rows, bom_table_cnt, skipped_tables


def main():
    if not SRC.exists():
        print(f"ERROR: {SRC} 不存在")
        sys.exit(1)
    print(f"Reading {SRC} ({SRC.stat().st_size / 1024:.0f} KB)…")
    doc = Document(SRC)
    print(f"Total tables in docx: {len(doc.tables)}")

    rows, bom_tbl, skip_tbl = extract_rows_from_doc(doc)
    print(f"BOM tables detected: {bom_tbl}; non-BOM tables skipped: {skip_tbl}")
    print(f"Rows extracted: {len(rows)}")

    df = pd.DataFrame(rows)
    # 去重：相同 part_id + nomenclature
    before = len(df)
    df = df.drop_duplicates(subset=["part_id", "nomenclature"], keep="first")
    print(f"Dedup: {before} → {len(df)}")

    # EFFECT=A 过滤（PT6A-60A 型号）
    # 规则：effect_code 包含 'A' 字母，或为空（通用零件）
    # 使用更宽松规则：保留所有
    # （可按需启用 EFFECT=A 过滤）
    # df = df[df['effect_code'].str.contains('A', na=True, regex=False) | (df['effect_code'].str.strip() == '')]

    # 类别统计
    print("\nCategory distribution:")
    print(df["category"].value_counts())

    DST.parent.mkdir(parents=True, exist_ok=True)
    df.to_csv(DST, index=False, encoding="utf-8-sig")
    print(f"\n[OK] Saved: {DST} ({DST.stat().st_size / 1024:.0f} KB)")

    # 预览首尾
    print("\n=== first 3 rows ===")
    print(df.head(3).to_string(index=False, max_colwidth=40))
    print("\n=== last 3 rows ===")
    print(df.tail(3).to_string(index=False, max_colwidth=40))


if __name__ == "__main__":
    main()
