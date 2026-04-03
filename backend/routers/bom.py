"""
backend/routers/bom.py — BOM 图数据库接口

GET  /bom/status    — Neo4j 连接状态（JSON）
POST /bom/ingest    — BOM 文件 → Neo4j 入库（支持 xlsx/pdf/docx，multipart + SSE）
POST /bom/query     — 关键词查询 BOM 图谱（JSON）

支持 PDF/DOCX 上传：先提取文本/表格，再调用 LLM 自动转换为标准 BOM 格式。
Neo4j 驱动懒加载：首次调用时建立连接，用 state.neo4j_lock 防止并发重复连接。
"""

import json
import os
import re
import tempfile
from typing import Optional

import pandas as pd
import pdfplumber
from fastapi import APIRouter, Depends, File, Form, Request, UploadFile
from pydantic import BaseModel

from backend.deps import get_state, get_neo4j_cfg
from backend.state import AppState
from backend.sse import log_gen_to_sse

try:
    from docx import Document as DocxDocument
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False

router = APIRouter(prefix="/bom")

try:
    from neo4j import GraphDatabase as _Neo4jDriver
    _NEO4J_AVAILABLE = True
except ImportError:
    _NEO4J_AVAILABLE = False


# ==========================================
# Neo4j 懒加载辅助函数
# ==========================================

def _get_neo4j_driver(state: AppState, cfg: dict):
    """获取 Neo4j 驱动，首次建立连接时加锁，失败返回 None。"""
    if state.neo4j_driver is not None:
        return state.neo4j_driver
    if not _NEO4J_AVAILABLE:
        return None
    with state.neo4j_lock:
        if state.neo4j_driver is not None:
            return state.neo4j_driver
        try:
            driver = _Neo4jDriver.driver(cfg["uri"], auth=(cfg["user"], cfg["pass"]))
            driver.verify_connectivity()
            state.neo4j_driver = driver
        except Exception:
            state.neo4j_driver = None
    return state.neo4j_driver


# ==========================================
# PDF / DOCX 文本提取
# ==========================================

def _extract_from_pdf(filepath: str) -> str:
    """从 PDF 提取文本，优先提取表格（转 Markdown table），无表格则提取纯文本。"""
    all_content = []
    with pdfplumber.open(filepath) as pdf:
        for i, page in enumerate(pdf.pages):
            tables = page.extract_tables()
            if tables:
                for table in tables:
                    if not table or not table[0]:
                        continue
                    header = table[0]
                    md = "| " + " | ".join(str(c or "") for c in header) + " |\n"
                    md += "| " + " | ".join("---" for _ in header) + " |\n"
                    for row in table[1:]:
                        md += "| " + " | ".join(str(c or "") for c in row) + " |\n"
                    all_content.append(f"[Page {i + 1} Table]\n{md}")
            else:
                text = page.extract_text() or ""
                if text.strip():
                    all_content.append(f"[Page {i + 1}]\n{text}")
    return "\n\n".join(all_content)


def _extract_from_docx(filepath: str) -> str:
    """从 DOCX 提取表格（转 Markdown table）和标题段落。
    BOM 数据主要在表格中，段落仅保留标题级别的作为上下文，避免正文过多。"""
    doc = DocxDocument(filepath)
    all_content = []
    # 仅保留标题段落作为上下文（Heading 1-4 或短段落）
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        is_heading = para.style and para.style.name and "Heading" in para.style.name
        if is_heading or (len(text) < 80 and not text.endswith("。")):
            all_content.append(text)
    # 提取所有表格（BOM 数据的主要来源）
    for idx, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        if rows:
            header = rows[0]
            md = "| " + " | ".join(header) + " |\n"
            md += "| " + " | ".join("---" for _ in header) + " |\n"
            for row in rows[1:]:
                md += "| " + " | ".join(row) + " |\n"
            all_content.append(f"[Table {idx + 1}]\n{md}")
    return "\n\n".join(all_content)


# ==========================================
# LLM BOM 格式转换
# ==========================================

BOM_CONVERSION_PROMPT = """你是一个BOM（物料清单）数据结构化专家。请将以下文档内容解析为标准BOM格式的JSON数组。

## 必填字段说明
- level_code: 层级编码，如 "1", "1.1", "1.1.1"（用.分隔，表示父子关系层级）
- part_id: 零件编号/图号（如文档中无编号，请根据零件名称生成如 P001, P002 格式的编号）
- part_name: 零件中文名称
- part_name_en: 零件英文名称（如无法确定，留空字符串）
- category: 分类，必须是以下三者之一：
  - "Assembly"：组件/总成（包含子零件的装配体）
  - "Part"：零件（最底层不可拆分的零件）
  - "Standard"：标准件（螺栓、螺母、垫圈、轴承等通用标准件）
- qty: 数量（整数，默认1）
- unit: 单位（件、个、套等，默认"件"）

## 选填字段
- material: 材料（如钛合金、不锈钢等，无则留空字符串）
- weight_kg: 重量(kg)，数字字符串或空字符串
- spec: 规格型号（无则留空字符串）
- note: 备注（无则留空字符串）

## 层级编码规则
- 最顶层组件编码为 "1"（如有多个顶层则为 "1", "2", "3"...）
- 子零件编码在父零件后追加，如 "1.1", "1.2"
- 更深层级类推："1.1.1", "1.1.2"
- 根据文档中的缩进、编号或上下文关系推断层级

## 输出要求
仅输出JSON数组，不要有任何其他说明文字。示例：
```json
[
  {{"level_code":"1","part_id":"ASM-001","part_name":"高压压气机","part_name_en":"HPC","category":"Assembly","qty":1,"unit":"套","material":"","weight_kg":"","spec":"","note":""}},
  {{"level_code":"1.1","part_id":"BLD-001","part_name":"第一级叶片","part_name_en":"Stage 1 Blade","category":"Part","qty":36,"unit":"件","material":"钛合金TC4","weight_kg":"0.5","spec":"","note":""}}
]
```

## 待解析的文档内容：
{content}"""


def _split_for_llm(text: str, max_chars: int = 15000) -> list:
    """按 [Page N] / [Table N] 标记分割文本，保证单个表格不被截断。
    max_chars 默认 15000（约 5000-7000 中文 token），平衡调用次数与上下文质量。"""
    sections = re.split(r"(?=\[(?:Page|Table) \d+)", text)
    chunks, current = [], ""
    for section in sections:
        if len(current) + len(section) > max_chars and current:
            chunks.append(current)
            current = section
        else:
            current += section
    if current:
        chunks.append(current)
    return chunks if chunks else [text]


def _parse_llm_json(raw: str) -> list:
    """从 LLM 输出中提取 JSON 数组，容错处理。"""
    raw = re.sub(r"```json\s*", "", raw)
    raw = re.sub(r"```\s*", "", raw)
    raw = raw.strip()
    try:
        data = json.loads(raw)
        return data if isinstance(data, list) else [data]
    except json.JSONDecodeError:
        pass
    match = re.search(r"\[.*\]", raw, re.DOTALL)
    if match:
        try:
            return json.loads(match.group())
        except json.JSONDecodeError:
            pass
    return []


class _LlmBomConverter:
    """调用 LLM 将提取的原始文本分段转换为标准 BOM DataFrame。
    用 run() 生成器逐段 yield 日志消息，转换完成后从 .df 获取结果。"""

    def __init__(self, raw_text: str, state: AppState):
        self._raw_text = raw_text
        self._state = state
        self.df: Optional[pd.DataFrame] = None

    def run(self):
        """生成器：逐段调用 LLM，每一步 yield 一条日志字符串。"""
        chunks = _split_for_llm(self._raw_text)
        yield f"   文档分为 {len(chunks)} 段，开始 LLM 转换…"

        all_records = []
        for i, chunk in enumerate(chunks):
            yield f"   🤖 LLM 转换第 {i + 1}/{len(chunks)} 段…"
            prompt = BOM_CONVERSION_PROMPT.format(content=chunk)
            try:
                resp = self._state.llm_client.chat.completions.create(
                    model=None,
                    messages=[{"role": "user", "content": prompt}],
                    temperature=0.1,
                )
                raw_json = resp.choices[0].message.content
                records = _parse_llm_json(raw_json)
                yield f"   ✅ 第 {i + 1} 段提取到 {len(records)} 条记录"
                all_records.extend(records)
            except Exception as e:
                yield f"   ⚠️ 第 {i + 1} 段 LLM 调用失败：{e}"

        if not all_records:
            self.df = None
            return

        df = pd.DataFrame(all_records)
        required = ["level_code", "part_id", "part_name", "part_name_en",
                     "category", "qty", "unit"]
        optional = ["material", "weight_kg", "spec", "note"]
        for col in required + optional:
            if col not in df.columns:
                df[col] = ""
        df["qty"] = pd.to_numeric(df["qty"], errors="coerce").fillna(1).astype(int)
        df["unit"] = df["unit"].replace("", "件")
        df["category"] = df["category"].replace("", "Part")
        df = df.dropna(subset=["level_code", "part_id"])
        df = df[df["level_code"].astype(str).str.strip() != ""]
        df = df[df["part_id"].astype(str).str.strip() != ""]
        for col in df.columns:
            df[col] = df[col].fillna("").astype(str).str.strip()
        df["qty"] = pd.to_numeric(df["qty"], errors="coerce").fillna(1).astype(int)

        self.df = df
        yield f"   ✅ LLM 转换完成：共 {len(df)} 条有效记录"


# ==========================================
# GET /bom/status
# ==========================================

@router.get("/status", summary="Neo4j 连接状态")
def bom_status(state: AppState = Depends(get_state), cfg: dict = Depends(get_neo4j_cfg)):
    """检测 Neo4j 连接，返回节点/关系数量。"""
    # 强制重新连接（等同于 app.py 中的 _neo4j_driver = None 重置）
    with state.neo4j_lock:
        if state.neo4j_driver is not None:
            try:
                state.neo4j_driver.close()
            except Exception:
                pass
        state.neo4j_driver = None

    if not _NEO4J_AVAILABLE:
        return {"connected": False, "error": "neo4j 包未安装，请运行 pip install neo4j", "uri": cfg["uri"]}

    driver = _get_neo4j_driver(state, cfg)
    if driver is None:
        return {"connected": False, "error": "无法连接 Neo4j，请确认 Docker 容器已启动", "uri": cfg["uri"]}

    try:
        with driver.session() as s:
            nodes = s.run("""
                MATCH (n) WHERE n:Assembly OR n:Part OR n:Standard
                RETURN count(n) AS cnt
            """).single()["cnt"]
            edges = s.run("MATCH ()-[r:CHILD_OF]->() RETURN count(r) AS cnt").single()["cnt"]
        return {"connected": True, "nodes": nodes, "edges": edges, "uri": cfg["uri"]}
    except Exception as e:
        return {"connected": False, "error": str(e), "uri": cfg["uri"]}


# ==========================================
# POST /bom/ingest
# ==========================================

@router.post("/ingest/pipeline", summary="LangGraph 管道 BOM/CAD 入库（SSE）")
def bom_ingest_pipeline(
    request: Request,
    file: Optional[UploadFile] = File(default=None),
    clear_first: bool = Form(default=False),
    state: AppState = Depends(get_state),
    cfg: dict = Depends(get_neo4j_cfg),
):
    """
    LangGraph 管道入库（BOM 与 CAD 统一入口）：
    - STEP/STP → lg_cad_pipeline（parse_cad_step → cad_to_kg_triples）
    - PDF/DOCX  → lg_bom_pipeline（extract_tables → llm_to_csv → write_neo4j）
    - Excel/CSV → lg_bom_pipeline（load_table → validate_bom_df → write_neo4j）
    两条管道均在后台线程运行（同 RAG 管道），LLM 调用期间实时推送进度日志。
    响应为 SSE 日志流。
    """
    import queue
    from backend.pipelines.sse_bridge import pipeline_to_log_generator

    bom_default = request.app.state.bom_default

    if file is None:
        filepath = bom_default
        filename = os.path.basename(filepath)
        tmp_path = None
    else:
        suffix = os.path.splitext(file.filename or "bom.xlsx")[1] or ".xlsx"
        with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
            tmp.write(file.file.read())
            tmp_path = tmp.name
        filepath = tmp_path
        filename = file.filename

    if not os.path.exists(filepath):
        def _err():
            yield f"❌ 文件不存在：{filepath}"
        return log_gen_to_sse(_err())

    # 根据扩展名选择管道
    ext = os.path.splitext(filename or "")[1].lower()
    is_cad = ext in (".step", ".stp")

    if is_cad:
        pipeline = state.lg_cad_pipeline
        mode = "cad"
        if pipeline is None:
            def _err():
                yield "❌ LangGraph CAD 管道未初始化，请检查后端日志"
            return log_gen_to_sse(_err())
    else:
        pipeline = state.lg_bom_pipeline
        mode = "bom"
        if pipeline is None:
            def _err():
                yield "❌ LangGraph BOM 管道未初始化，请检查后端日志"
            return log_gen_to_sse(_err())

    initial_state = {
        "file_path": filepath,
        "pipeline_mode": mode,
        "clear_first": clear_first,
        "log_messages": [f"[pipeline] 开始处理（{mode.upper()}）：{filename}"],
    }

    # 使用进度队列 + 后台线程，避免 LLM 调用期间 SSE 静默挂起
    # 同时挂到 state._ingest_progress_q，让节点内部（如 KG 节点）可通过侧信道推送中间进度
    progress_q = queue.Queue()
    state._ingest_progress_q = progress_q

    def _cleanup_gen():
        try:
            yield from pipeline_to_log_generator(pipeline, initial_state,
                                                  progress_queue=progress_q)
        finally:
            state._ingest_progress_q = None  # 清理侧信道
            if tmp_path and os.path.exists(tmp_path):
                try:
                    os.remove(tmp_path)
                except Exception:
                    pass

    return log_gen_to_sse(_cleanup_gen())


@router.post("/ingest", summary="BOM 文件入库（xlsx/pdf/docx，SSE）")
def bom_ingest(
    request: Request,
    file: Optional[UploadFile] = File(default=None),
    clear_first: bool = Form(default=False),
    state: AppState = Depends(get_state),
    cfg: dict = Depends(get_neo4j_cfg),
):
    """
    上传 BOM 文件（xlsx/pdf/docx，可选），写入 Neo4j。
    PDF/DOCX 会自动通过 LLM 转换为标准 BOM 格式。
    响应为 SSE 日志流。
    """
    bom_default = request.app.state.bom_default
    gen = _run_bom_ingest(state, cfg, file, clear_first, bom_default)
    return log_gen_to_sse(gen)


def _run_bom_ingest(state: AppState, cfg: dict, file, clear_first: bool, bom_default: str):
    """BOM 入库生成器，迁移自 app.py run_bom_ingest_ui()。"""
    lines = []

    def emit(msg):
        lines.append(msg)
        return "\n".join(lines)

    tmp_path = None
    try:
        # 确定文件路径
        if file is None:
            filepath = bom_default
            yield emit(f"📌 未上传文件，使用默认测试 BOM：{os.path.basename(filepath)}")
        else:
            # 将上传文件写入临时路径
            suffix = os.path.splitext(file.filename or "bom.xlsx")[1] or ".xlsx"
            with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
                tmp.write(file.file.read())
                tmp_path = tmp.name
            filepath = tmp_path
            yield emit(f"📌 使用上传文件：{file.filename}")

        if not os.path.exists(filepath):
            yield emit(f"❌ 文件不存在：{filepath}")
            return

        # 根据文件类型读取 BOM
        ext = os.path.splitext(filepath)[1].lower()
        yield emit(f"📖 读取 BOM 表（{ext}）…")

        try:
            if ext in (".xlsx", ".xls"):
                # Excel 直接读取
                df = pd.read_excel(filepath, dtype=str)
                df.columns = [c.strip().lower() for c in df.columns]
                for col in ("material", "weight_kg", "spec", "note"):
                    if col not in df.columns:
                        df[col] = ""
                df = df.dropna(subset=["level_code", "part_id"])
                for col in df.columns:
                    df[col] = df[col].fillna("").astype(str).str.strip()
                df["qty"] = pd.to_numeric(df["qty"], errors="coerce").fillna(1).astype(int)

            elif ext == ".pdf":
                yield emit("📄 正在提取 PDF 内容…")
                raw_text = _extract_from_pdf(filepath)
                if not raw_text.strip():
                    yield emit("❌ PDF 中未找到可识别的文本或表格")
                    return
                yield emit(f"✅ PDF 提取完成（{len(raw_text)} 字符）")
                yield emit("🤖 调用 LLM 自动转换 BOM 格式…")
                converter = _LlmBomConverter(raw_text, state)
                for log_line in converter.run():
                    yield emit(log_line)
                df = converter.df
                if df is None or df.empty:
                    yield emit("❌ LLM 未能从文档中提取有效 BOM 记录，请检查文档内容")
                    return

            elif ext == ".docx":
                if not _DOCX_AVAILABLE:
                    yield emit("❌ 缺少 python-docx，请运行：pip install python-docx")
                    return
                yield emit("📄 正在提取 DOCX 内容…")
                raw_text = _extract_from_docx(filepath)
                if not raw_text.strip():
                    yield emit("❌ DOCX 中未找到可识别的文本或表格")
                    return
                yield emit(f"✅ DOCX 提取完成（{len(raw_text)} 字符）")
                yield emit("🤖 调用 LLM 自动转换 BOM 格式…")
                converter = _LlmBomConverter(raw_text, state)
                for log_line in converter.run():
                    yield emit(log_line)
                df = converter.df
                if df is None or df.empty:
                    yield emit("❌ LLM 未能从文档中提取有效 BOM 记录，请检查文档内容")
                    return

            elif ext == ".doc":
                yield emit("❌ 不支持旧版 .doc 格式，请将文件另存为 .docx 后重新上传")
                return
            else:
                yield emit(f"❌ 不支持的文件格式：{ext}（支持 xlsx/pdf/docx）")
                return

            yield emit(f"✅ 读取成功：共 {len(df)} 条零件记录")
        except Exception as e:
            yield emit(f"❌ 读取失败：{e}")
            return

        # 连接 Neo4j
        yield emit(f"🔌 连接 Neo4j（{cfg['uri']}）…")
        driver = _get_neo4j_driver(state, cfg)
        if driver is None:
            yield emit("❌ Neo4j 不可用，请先启动：docker start neo4j")
            return
        yield emit("✅ 连接成功")

        with driver.session() as session:
            for label in ("Assembly", "Part", "Standard"):
                session.run(f"CREATE CONSTRAINT IF NOT EXISTS FOR (n:{label}) REQUIRE n.part_id IS UNIQUE")
            yield emit("🔧 Schema 约束已就绪")

            if clear_first:
                r = session.run("""
                    MATCH (n) WHERE n:Assembly OR n:Part OR n:Standard
                    DETACH DELETE n RETURN count(n) AS cnt
                """).single()
                yield emit(f"🗑️ 已清空旧数据（{r['cnt']} 个节点）")

            records = df.to_dict(orient="records")
            for category in ("Assembly", "Part", "Standard"):
                batch = [r for r in records if r["category"] == category]
                if not batch:
                    continue
                session.run(f"""
                    UNWIND $rows AS row
                    MERGE (n:{category} {{part_id: row.part_id}})
                    SET n.level_code   = row.level_code,
                        n.part_name    = row.part_name,
                        n.part_name_en = row.part_name_en,
                        n.qty          = toInteger(row.qty),
                        n.unit         = row.unit,
                        n.material     = row.material,
                        n.weight_kg    = CASE row.weight_kg WHEN '' THEN null ELSE toFloat(row.weight_kg) END,
                        n.spec         = row.spec,
                        n.note         = row.note
                """, rows=batch)
                yield emit(f"   ✅ {category}：写入 {len(batch)} 个节点")

            level_map = dict(zip(df["level_code"], df["part_id"]))
            edges = []
            for _, row in df.iterrows():
                code = str(row["level_code"])
                if "." not in code:
                    continue
                parent_code = ".".join(code.split(".")[:-1])
                parent_id = level_map.get(parent_code)
                if parent_id:
                    edges.append({"child_id": row["part_id"], "parent_id": parent_id})
            if edges:
                session.run("""
                    UNWIND $edges AS e
                    MATCH (child  {part_id: e.child_id})
                    MATCH (parent {part_id: e.parent_id})
                    MERGE (child)-[:CHILD_OF]->(parent)
                """, edges=edges)
            yield emit(f"   ✅ CHILD_OF 关系：建立 {len(edges)} 条")

            r2 = session.run("""
                MATCH (n) WHERE n:Assembly OR n:Part OR n:Standard
                RETURN count(CASE WHEN n:Assembly THEN 1 END) AS a,
                       count(CASE WHEN n:Part     THEN 1 END) AS p,
                       count(CASE WHEN n:Standard THEN 1 END) AS s
            """).single()
            yield emit(
                f"\n🎉 入库完成！Assembly:{r2['a']}  Part:{r2['p']}  Standard:{r2['s']}  "
                f"关系:{len(edges)} 条\n可视化查看：http://localhost:7474"
            )
    finally:
        if tmp_path and os.path.exists(tmp_path):
            try:
                os.remove(tmp_path)
            except Exception:
                pass


# ==========================================
# POST /bom/query
# ==========================================

class BomQueryRequest(BaseModel):
    question: str


@router.post("/query", summary="查询 BOM 图谱")
def bom_query(body: BomQueryRequest, state: AppState = Depends(get_state), cfg: dict = Depends(get_neo4j_cfg)):
    """从 Neo4j 查询与问题相关的 BOM 零件信息，返回格式化文本。"""
    bom_text = _query_bom_text(state, cfg, body.question)
    return {"bom_text": bom_text}


# ==========================================
# GET /bom/kg/graph — 知识图谱可视化数据
# ==========================================

@router.get("/kg/graph", summary="获取知识图谱可视化数据（D3 力导向图）")
def kg_graph(
    keyword: str = "",
    limit: int = 200,
    state: AppState = Depends(get_state),
    cfg: dict = Depends(get_neo4j_cfg),
):
    """
    返回 Neo4j 中所有类型的节点和关系，格式化为 D3 力导向图数据。

    - keyword 为空：返回全图（截断到 limit）
    - keyword 非空：返回包含该关键词的节点及其 1 跳邻居
    """
    driver = _get_neo4j_driver(state, cfg)
    if driver is None:
        return {"nodes": [], "links": [], "error": "Neo4j 未连接"}

    nodes_map: dict = {}   # id → node dict
    links: list = []
    seen_links: set = set()

    def _node_id(n) -> str:
        """优先使用 kg_id，其次 part_id，兜底 elementId。"""
        return n.get("kg_id") or n.get("part_id") or str(n.element_id)

    def _node_label(n) -> str:
        return n.get("kg_name") or n.get("part_name") or "?"

    def _add_node(n):
        nid = _node_id(n)
        if nid in nodes_map:
            return nid
        labels_list = list(n.labels) if hasattr(n, 'labels') else []
        node_type = labels_list[0] if labels_list else "Unknown"
        nodes_map[nid] = {
            "id": nid,
            "label": _node_label(n),
            "type": node_type,
        }
        # 附加有用属性
        for key in ("part_id", "kg_name", "part_name", "ata_section", "material",
                     "spec", "qty", "unit", "weight_kg", "seq_no"):
            val = n.get(key)
            if val is not None:
                nodes_map[nid][key] = val
        return nid

    def _add_link(src_id: str, tgt_id: str, rel_type: str):
        key = (src_id, rel_type, tgt_id)
        if key not in seen_links:
            seen_links.add(key)
            links.append({"source": src_id, "target": tgt_id, "type": rel_type})

    with driver.session() as session:
        if keyword.strip():
            # 关键词模式：匹配节点 + 1 跳邻居
            rows = session.run("""
                MATCH (n)
                WHERE n.kg_name CONTAINS $kw
                   OR n.part_name CONTAINS $kw
                   OR (n.ata_section IS NOT NULL AND n.ata_section CONTAINS $kw)
                WITH n LIMIT $lim
                OPTIONAL MATCH (n)-[r]-(m)
                RETURN n, r, m
            """, kw=keyword.strip(), lim=limit).data()
        else:
            # 全图模式
            rows = session.run("""
                MATCH (n)
                WITH n LIMIT $lim
                OPTIONAL MATCH (n)-[r]->(m)
                RETURN n, r, m
            """, lim=limit).data()

        for row in rows:
            n = row.get("n")
            if n is not None:
                _add_node(n)

            m = row.get("m")
            r = row.get("r")
            if m is not None and r is not None:
                mid = _add_node(m)
                nid = _node_id(n)
                rel_type = r.type if hasattr(r, 'type') else str(type(r))
                _add_link(nid, mid, rel_type)

    return {"nodes": list(nodes_map.values()), "links": links}


def _topological_sort_kahn(nodes: list, edges: list, node_data: dict) -> list:
    """
    Kahn 算法拓扑排序，返回有序 node_data 列表。
    孤立节点（无 precedes 关系）按 seq_no 追加到末尾。
    """
    from collections import deque, defaultdict

    in_degree: dict = defaultdict(int)
    adj: dict = defaultdict(list)
    node_set = set(nodes)

    for u, v in edges:
        if u in node_set and v in node_set:
            adj[u].append(v)
            in_degree[v] += 1

    def _seq_key(nid: str) -> float:
        d = node_data.get(nid, {})
        s = d.get("seq_no")
        try:
            return float(s) if s is not None else 9999.0
        except (TypeError, ValueError):
            return 9999.0

    queue = deque(
        sorted([n for n in node_set if in_degree[n] == 0], key=_seq_key)
    )
    result: list = []
    sorted_node_ids: set = set()
    while queue:
        node = queue.popleft()
        result.append(node_data[node])
        sorted_node_ids.add(node)
        for neighbor in sorted(adj[node], key=_seq_key):
            in_degree[neighbor] -= 1
            if in_degree[neighbor] == 0:
                queue.append(neighbor)

    # 追加孤立节点（有环或无 precedes 边的节点）
    for n in sorted(node_set - sorted_node_ids, key=_seq_key):
        result.append(node_data[n])
    return result


def _query_procedure_chain(
    state: AppState, cfg: dict, question: str
) -> tuple[list, str]:
    """
    查询 precedes 关系链，返回 (有序步骤列表, 格式化文本)。

    流程：
      1. 按关键词查 Procedure 种子节点
      2. 沿 precedes 链扩展（最多 5 跳）
      3. 拉取 Tool / Specification 信息
      4. Python 层 Kahn 拓扑排序
      5. 格式化输出

    Returns:
        (ordered_steps, formatted_text)
    """
    driver = _get_neo4j_driver(state, cfg)
    if driver is None:
        return [], ""

    kw = question[:20]

    with driver.session() as session:
        cnt = session.run(
            "MATCH (p:Procedure) RETURN count(p) AS cnt"
        ).single()
        if not cnt or cnt["cnt"] == 0:
            return [], ""

        seed_rows = session.run("""
            MATCH (p:Procedure)
            WHERE p.kg_name CONTAINS $kw
               OR (p.ata_section IS NOT NULL AND p.ata_section CONTAINS $kw)
            RETURN p.kg_id AS proc_id, p.kg_name AS proc_name,
                   p.ata_section AS ata_section, p.seq_no AS seq_no
            LIMIT 20
        """, kw=kw).data()

        if not seed_rows:
            return [], ""

        seed_ids = [r["proc_id"] for r in seed_rows]

        chain_rows = session.run("""
            MATCH (start:Procedure)-[:precedes*0..5]->(n:Procedure)
            WHERE start.kg_id IN $seed_ids
            WITH DISTINCT n
            OPTIONAL MATCH (n)-[:requires]->(t:Tool)
            OPTIONAL MATCH (n)-[:specifiedBy]->(s:Specification)
            RETURN n.kg_id AS proc_id, n.kg_name AS proc_name,
                   n.seq_no AS seq_no, n.ata_section AS ata_section,
                   collect(DISTINCT t.kg_name) AS tools,
                   collect(DISTINCT (
                       COALESCE(s.spec_value, '') + COALESCE(s.spec_unit, '')
                   )) AS specs
        """, seed_ids=seed_ids).data()

        if not chain_rows:
            return [], ""

        all_ids = [r["proc_id"] for r in chain_rows]
        edge_rows = session.run("""
            MATCH (a:Procedure)-[:precedes]->(b:Procedure)
            WHERE a.kg_id IN $all_ids AND b.kg_id IN $all_ids
            RETURN a.kg_id AS from_id, b.kg_id AS to_id
        """, all_ids=all_ids).data()

    node_data = {r["proc_id"]: r for r in chain_rows}
    edges = [(e["from_id"], e["to_id"]) for e in edge_rows]
    ordered = _topological_sort_kahn(list(node_data.keys()), edges, node_data)

    lines = ["【装配工序（有序）】"]
    for i, step in enumerate(ordered, 1):
        tools = "、".join(t for t in (step.get("tools") or []) if t)
        specs = "、".join(s for s in (step.get("specs") or []) if s)
        line = f"  步骤{i}：{step['proc_name']}"
        if tools:
            line += f"  工具：{tools}"
        if specs:
            line += f"  规范：{specs}"
        lines.append(line)

    return ordered, "\n".join(lines)


def _query_bom_text(state: AppState, cfg: dict, question: str) -> str:
    """BOM 图谱查询，迁移自 app.py _query_bom_text()。"""
    driver = _get_neo4j_driver(state, cfg)
    if driver is None:
        return ""

    KNOWN_KEYWORDS = [
        "风扇", "压气机", "高压压气机", "低压压气机", "燃烧室",
        "高压涡轮", "低压涡轮", "涡轮", "附件", "尾喷管",
        "叶片", "涡轮盘", "火焰筒", "喷嘴", "机匣", "转子", "静子",
    ]
    keywords = [kw for kw in KNOWN_KEYWORDS if kw in question] or [question[:15]]

    with driver.session() as session:
        lines = []
        seen = set()
        for kw in keywords:
            rows = session.run("""
                MATCH (parent) WHERE parent.part_name CONTAINS $kw
                OPTIONAL MATCH (child)-[:CHILD_OF*1..2]->(parent)
                WITH parent, collect(child) AS children
                RETURN parent.part_name AS module,
                       parent.part_id   AS module_id,
                       parent.spec      AS spec,
                       [c IN children | {
                         name:c.part_name, id:c.part_id,
                         qty:c.qty, unit:c.unit,
                         material:c.material, spec:c.spec, note:c.note
                       }] AS parts
            """, kw=kw).data()

            for r in rows:
                mod = r.get("module", "")
                if not mod or mod in seen:
                    continue
                seen.add(mod)
                lines.append(f"【{mod}】({r.get('module_id','')})")
                if r.get("spec"):
                    lines.append(f"  规格：{r['spec']}")
                for p in (r.get("parts") or []):
                    mat  = f" 材料:{p['material']}" if p.get("material") else ""
                    spec = f" 规格:{p['spec']}"     if p.get("spec") else ""
                    note = f" 备注:{p['note']}"     if p.get("note") else ""
                    lines.append(
                        f"  - {p['name']}({p['id']}) ×{p['qty']}{p['unit']}{mat}{spec}{note}"
                    )
        return "\n".join(lines) if lines else ""


def _query_bom_entities(state: AppState, cfg: dict, question: str) -> dict:
    """
    BOM 图谱结构化查询：返回独立实体列表和关系列表，供细粒度溯源使用。
    每个 Neo4j 节点和每条 CHILD_OF 边分别作为独立 Citation。
    无 Neo4j 连接时返回 {"entities": [], "relations": []}，不抛异常。
    """
    driver = _get_neo4j_driver(state, cfg)
    if driver is None:
        return {"entities": [], "relations": []}

    KNOWN_KEYWORDS = [
        "风扇", "压气机", "高压压气机", "低压压气机", "燃烧室",
        "高压涡轮", "低压涡轮", "涡轮", "附件", "尾喷管",
        "叶片", "涡轮盘", "火焰筒", "喷嘴", "机匣", "转子", "静子",
    ]
    keywords = [kw for kw in KNOWN_KEYWORDS if kw in question] or [question[:15]]

    seen_entities: dict = {}      # part_id → entity dict（按 part_id 去重）
    seen_rel_keys: set = set()    # (from_id, rel_type, to_id)
    relations: list = []

    def _safe(val):
        """将 None 转为空字符串，保留数值类型。"""
        if val is None:
            return ""
        if isinstance(val, str):
            return val
        return val  # int / float 原样保留

    with driver.session() as session:
        for kw in keywords:
            # ── 查询 A：匹配节点本身（附带直接父节点信息）──────────
            for r in session.run("""
                MATCH (n) WHERE n.part_name CONTAINS $kw
                OPTIONAL MATCH (n)-[:CHILD_OF]->(parent)
                RETURN labels(n)[0]      AS entity_type,
                       n.part_id         AS part_id,
                       n.part_name       AS part_name,
                       n.part_name_en    AS part_name_en,
                       n.qty             AS qty,
                       n.unit            AS unit,
                       n.material        AS material,
                       n.weight_kg       AS weight_kg,
                       n.spec            AS spec,
                       n.note            AS note,
                       n.level_code      AS level_code,
                       parent.part_id    AS parent_id,
                       parent.part_name  AS parent_name
                LIMIT 30
            """, kw=kw).data():
                pid = r.get("part_id")
                if pid and pid not in seen_entities:
                    seen_entities[pid] = {
                        "entity_type":  _safe(r.get("entity_type")) or "Part",
                        "part_id":      _safe(r.get("part_id")),
                        "part_name":    _safe(r.get("part_name")),
                        "part_name_en": _safe(r.get("part_name_en")),
                        "qty":          r.get("qty"),
                        "unit":         _safe(r.get("unit")),
                        "material":     _safe(r.get("material")),
                        "weight_kg":    r.get("weight_kg"),
                        "spec":         _safe(r.get("spec")),
                        "note":         _safe(r.get("note")),
                        "level_code":   _safe(r.get("level_code")),
                        "parent_id":    r.get("parent_id"),
                        "parent_name":  r.get("parent_name"),
                    }

            # ── 查询 B：直接子节点（depth=1）──────────────────────
            for r in session.run("""
                MATCH (n) WHERE n.part_name CONTAINS $kw
                MATCH (child)-[:CHILD_OF]->(n)
                RETURN labels(child)[0]   AS entity_type,
                       child.part_id      AS part_id,
                       child.part_name    AS part_name,
                       child.part_name_en AS part_name_en,
                       child.qty          AS qty,
                       child.unit         AS unit,
                       child.material     AS material,
                       child.weight_kg    AS weight_kg,
                       child.spec         AS spec,
                       child.note         AS note,
                       child.level_code   AS level_code,
                       n.part_id          AS parent_id,
                       n.part_name        AS parent_name
                LIMIT 50
            """, kw=kw).data():
                pid = r.get("part_id")
                if pid and pid not in seen_entities:
                    seen_entities[pid] = {
                        "entity_type":  _safe(r.get("entity_type")) or "Part",
                        "part_id":      _safe(r.get("part_id")),
                        "part_name":    _safe(r.get("part_name")),
                        "part_name_en": _safe(r.get("part_name_en")),
                        "qty":          r.get("qty"),
                        "unit":         _safe(r.get("unit")),
                        "material":     _safe(r.get("material")),
                        "weight_kg":    r.get("weight_kg"),
                        "spec":         _safe(r.get("spec")),
                        "note":         _safe(r.get("note")),
                        "level_code":   _safe(r.get("level_code")),
                        "parent_id":    r.get("parent_id"),
                        "parent_name":  r.get("parent_name"),
                    }

            # ── 查询 C：CHILD_OF 关系边（去重）──────────────────────
            for r in session.run("""
                MATCH (n) WHERE n.part_name CONTAINS $kw
                MATCH (child)-[rel:CHILD_OF]->(n)
                RETURN type(rel)          AS rel_type,
                       labels(child)[0]   AS from_type,
                       child.part_id      AS from_id,
                       child.part_name    AS from_name,
                       labels(n)[0]       AS to_type,
                       n.part_id          AS to_id,
                       n.part_name        AS to_name
                LIMIT 60
            """, kw=kw).data():
                key = (r.get("from_id"), r.get("rel_type"), r.get("to_id"))
                if all(key) and key not in seen_rel_keys:
                    seen_rel_keys.add(key)
                    relations.append({
                        "rel_type":  _safe(r.get("rel_type")),
                        "from_type": _safe(r.get("from_type")),
                        "from_id":   _safe(r.get("from_id")),
                        "from_name": _safe(r.get("from_name")),
                        "to_type":   _safe(r.get("to_type")),
                        "to_id":     _safe(r.get("to_id")),
                        "to_name":   _safe(r.get("to_name")),
                    })

    return {"entities": list(seen_entities.values()), "relations": relations}
